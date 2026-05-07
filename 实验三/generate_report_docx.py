from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
ASSETS = BASE_DIR / "实验素材"
OUTPUT = BASE_DIR / "实验文档" / "2315302125 崔子霖3.docx"


def set_run_font(run, size=12, bold=False, east_asia="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_heading(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, east_asia="黑体")


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, east_asia="宋体", western="Courier New")


def add_picture(document, image_path, caption, width_cm):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = document.add_paragraph()
    format_paragraph(caption_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=11)


def build_cover(document):
    for _ in range(4):
        document.add_paragraph()

    school = document.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("南京理工大学紫金学院")
    set_run_font(run, size=22, bold=True, east_asia="黑体")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(42)
    run = title.add_run("《移动互联网开发基础》实验报告")
    set_run_font(run, size=20, bold=True, east_asia="黑体")

    table = document.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(8.8)

    fields = [
        ("课程编号：", "待补充"),
        ("课程名称：", "移动互联网开发基础"),
        ("学    院：", "计算机与人工智能"),
        ("专    业：", "计算机科学与技术"),
        ("学    号：", "2315302125"),
        ("姓    名：", "崔子霖"),
        ("任课教师：", "待补充"),
    ]

    for row, (label, value) in zip(table.rows, fields):
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        r2 = p2.add_run(value)
        set_run_font(r1, size=14)
        set_run_font(r2, size=14)

    for _ in range(8):
        document.add_paragraph()

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 5 月")
    set_run_font(run, size=14)

    document.add_page_break()


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    build_cover(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("实验三 装修网站首页")
    set_run_font(run, size=16, bold=True, east_asia="黑体")

    add_heading(document, "一、实验简介")
    add_body_paragraph(document, "本实验围绕装修网站首页设计展开，主要练习 HTML 基本结构、CSS 样式设置、网页布局、背景图像、导航栏、页眉和页脚等常见网页元素的实现方法。本次实验最终完成了一个装修公司首页，页面包含品牌导航、横幅背景图、整装服务、精选案例、装修流程和页脚联系信息。")
    add_body_paragraph(document, "本实验使用 macOS 环境完成，开发工具为 Visual Studio Code，运行浏览器为 Google Chrome。页面采用 HTML5 语义化结构和外部 CSS 样式表实现，并将图片素材统一放入 img 文件夹，符合实验指导书中关于图片文件夹和网页结构的要求。")

    add_heading(document, "二、实验要求")
    reqs = [
        "掌握 HTML 的基本结构和使用方法。",
        "掌握 CSS 的基本语法和样式设置方法。",
        "学会使用 HTML 和 CSS 进行网页布局。",
        "学会使用 CSS 设置背景图像和颜色。",
        "掌握网页中导航栏、页眉、内容区域和页脚等常见元素的样式设置。",
        "完成一个装修网站首页，并提供源代码和页面效果截图。"
    ]
    for index, item in enumerate(reqs, start=1):
        add_body_paragraph(document, f"{index}. {item}")

    add_heading(document, "三、概要设计")
    add_body_paragraph(document, "本实验页面采用典型企业网站首页结构，整体分为页眉导航区、横幅展示区、服务介绍区、案例展示区、装修流程区和页脚信息区。页眉区域放置网站 Logo 和导航菜单；横幅区域使用背景图片展示装修网站主题；服务区域以四个模块介绍装修服务；案例区域通过图片展示不同空间的装修效果；页脚区域展示公司简介、链接和联系方式。")
    add_body_paragraph(document, "网页文件结构分为 index.html、css/style.css、img 图片目录、实验文档目录和实验素材目录。HTML 负责页面内容结构，CSS 负责视觉样式和布局控制，图片文件负责页面背景和案例展示。")

    add_heading(document, "四、详细设计")
    add_body_paragraph(document, "HTML 部分使用 header、nav、main、section、article 和 footer 等语义化标签组织页面内容。导航栏包含“首页、整装服务、案例作品、装修流程、联系我们”等栏目；横幅区域放置装修网站主要宣传语；内容区域分别展示服务、案例和流程。HTML 结构代码截图如图 1 所示。")
    add_picture(document, ASSETS / "code-html.png", "图 1 HTML 页面结构代码截图", 15.5)
    add_body_paragraph(document, "CSS 部分通过外部样式表 css/style.css 完成页面样式设置。横幅区域使用 background-image 设置背景图片，并叠加 linear-gradient 提高文字可读性；页面布局主要使用 flex 和 grid 完成；同时使用媒体查询适配移动端显示。CSS 代码截图如图 2 所示。")
    add_picture(document, ASSETS / "code-css.png", "图 2 CSS 布局与背景图代码截图", 15.5)
    add_body_paragraph(document, "核心背景图样式如下：")
    add_code_block(document, """.banner {
  min-height: 620px;
  display: flex;
  align-items: center;
  background-image:
    linear-gradient(90deg, rgba(17, 39, 39, 0.82), rgba(17, 39, 39, 0.45)),
    url("../img/banner.jpg");
  background-repeat: no-repeat;
  background-position: center;
  background-size: cover;
}""")

    add_heading(document, "五、实验结果")
    add_body_paragraph(document, "本实验运行环境为 macOS（Apple M1 Pro），使用 Visual Studio Code 编写代码，使用 Google Chrome 打开本地 HTML 文件进行测试。实验三源码位于实验三目录下的 index.html 文件，CSS 样式文件位于 css/style.css，图片素材位于 img 文件夹。")
    steps = [
        "创建实验三目录，并新建 index.html、css/style.css、img、实验文档和实验素材等文件夹。",
        "编写 HTML 基本结构，完成页眉、导航栏、横幅、服务、案例、流程和页脚等模块。",
        "编写 CSS 样式，设置网页宽度、背景图片、颜色、网格布局、导航栏和页脚样式。",
        "将装修网站用到的图片放入 img 文件夹，并通过相对路径引用。",
        "使用 Chrome 运行页面，并分别生成桌面端和移动端页面效果截图。"
    ]
    for index, item in enumerate(steps, start=1):
        add_body_paragraph(document, f"{index}. {item}")
    add_body_paragraph(document, "图 3 为装修网站首页桌面端运行效果，图 4 为移动端运行效果。可以看到页面能够正常显示页眉导航、背景横幅、服务模块、案例图片和页面内容。")
    add_picture(document, ASSETS / "page-home.png", "图 3 装修网站首页桌面端效果图", 15.8)
    add_picture(document, ASSETS / "page-mobile.png", "图 4 装修网站首页移动端效果图", 7.2)
    add_body_paragraph(document, "测试结果表明，页面结构完整，图片素材能够正常加载，背景图像与颜色样式显示正确，导航栏、内容区域和页脚等常见网页元素均已实现，满足实验要求。")

    add_heading(document, "六、实验总结")
    add_body_paragraph(document, "通过本次实验，我进一步熟悉了 HTML 页面结构和 CSS 样式编写方法。相比前两次实验，本次实验更加侧重页面静态布局和视觉呈现，需要综合考虑导航栏、背景图片、内容模块、图片案例和页脚之间的布局关系。")
    add_body_paragraph(document, "实验过程中，重点在于合理拆分 HTML 结构和 CSS 样式，并通过相对路径正确引用图片文件。横幅背景图使用 CSS 的 background-image 属性实现，同时叠加渐变遮罩，使文字在图片上仍然保持清晰可读。")
    add_body_paragraph(document, "总体来看，本次实验完成了装修网站首页的设计与实现，掌握了企业类网站首页的基本组成方式，也提升了我对 HTML 语义结构、CSS 布局和响应式页面设计的理解。")

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
