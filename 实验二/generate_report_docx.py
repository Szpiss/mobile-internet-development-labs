from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
ASSETS = BASE_DIR / "实验素材"
OUTPUT = BASE_DIR / "实验文档" / "2315302125 崔子霖2.docx"


def set_run_font(run, size=12, bold=False, east_asia="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_heading(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, east_asia="黑体")


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, east_asia="宋体", western="Courier New")


def add_picture(document, image_path, caption, width_cm):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = document.add_paragraph()
    format_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    run = cap.add_run(caption)
    set_run_font(run, size=11)


def build_cover(document):
    for _ in range(4):
        document.add_paragraph()

    school = document.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("南京理工大学紫金学院")
    set_run_font(run, size=22, bold=True, east_asia="黑体")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(42)
    run = title.add_run("《移动互联网开发基础》实验报告")
    set_run_font(run, size=20, bold=True, east_asia="黑体")

    table = document.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(8.8)

    fields = [
        ("课程编号：", "待补充"),
        ("课程名称：", "移动互联网开发基础"),
        ("学    院：", "计算机与人工智能"),
        ("专    业：", "计算机科学与技术"),
        ("学    号：", "2315302125"),
        ("姓    名：", "崔子霖"),
        ("任课教师：", "待补充"),
    ]

    for row, (label, value) in zip(table.rows, fields):
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        r2 = p2.add_run(value)
        set_run_font(r1, size=14)
        set_run_font(r2, size=14)

    for _ in range(8):
        document.add_paragraph()

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 4 月")
    set_run_font(run, size=14)

    document.add_page_break()


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    build_cover(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("实验二 JavaScript 事件和表单页面设计")
    set_run_font(run, size=16, bold=True, east_asia="黑体")

    add_heading(document, "一、实验简介")
    add_body_paragraph(document, "本实验围绕 JavaScript 事件处理和 H5 表单页面设计展开，重点练习网页交互行为控制、表单元素使用方法以及前端数据合法性校验。实验内容分为两个部分：第一部分通过单击超链接改变正文的显示字号；第二部分通过 JavaScript 对用户登录表单中的用户名和密码进行合法性校验，并在输入不合法时给出明确错误提示。")
    add_body_paragraph(document, "由于本次实验使用的设备为 MacBook Pro（Apple M1 Pro），实验开发环境采用 macOS 平台下的 Visual Studio Code 和 Google Chrome。整个实验页面使用 HTML、CSS 和 JavaScript 编写，实现了窗口事件、鼠标事件、键盘事件、表单提交事件和表单重置事件等多种前端交互效果。")

    add_heading(document, "二、实验要求")
    reqs = [
        "理解 JavaScript 程序的概念与作用。",
        "理解事件发生时事件处理的三种方式。",
        "学会利用表单的提交及重置事件对表单数据进行校验。",
        "理解鼠标单击、鼠标移动、键盘输入和窗口加载等常见事件。",
        "掌握 H5 表单、输入框、按钮和字段集的基本语法与用法。",
        "完成字号切换和登录合法性校验两个实验项目。"
    ]
    for i, item in enumerate(reqs, start=1):
        add_body_paragraph(document, f"{i}. {item}")

    add_heading(document, "三、概要设计")
    add_body_paragraph(document, "本实验页面采用双栏布局，左侧用于展示“使用 JavaScript 改变网页字号大小”项目，右侧用于展示“用户登录合法性校验”项目。左侧区域提供“小”“中”“大”三个超链接，通过事件触发修改正文段落字号；右侧区域提供登录表单，包含用户名输入框、密码输入框、提交按钮和重置按钮，通过 JavaScript 对输入数据进行合法性判断。")
    add_body_paragraph(document, "在事件设计方面，页面同时使用了三种常见的事件绑定方式：行内事件、DOM0 事件和 addEventListener 事件。为增强演示效果，页面还增加了鼠标移动坐标显示、键盘输入状态显示和窗口加载状态显示。表单校验逻辑流程如图 1 所示。")
    add_picture(document, ASSETS / "validation-flow.png", "图 1 表单校验流程图", 14.2)

    add_heading(document, "四、详细设计")
    add_body_paragraph(document, "在项目 1 中，页面定义了 changeFontSize(size) 函数，根据传入的字号类型修改正文段落的 font-size，并同步更新当前字号状态。为了体现事件处理方式的区别，本实验将“小”字号链接设置为行内事件，将“中”字号链接设置为 DOM0 事件，将“大”字号链接设置为 addEventListener 事件。代码截图如图 2 所示。")
    add_picture(document, ASSETS / "code-font.png", "图 2 字号切换与事件绑定代码截图", 15.5)
    add_body_paragraph(document, "在项目 2 中，页面定义 validateLoginForm() 函数，对用户名和密码分别进行非空校验和长度校验，并在提交事件中先阻止表单默认提交，再根据校验结果显示错误提示或成功提示。当数据不合法时，页面会将焦点自动定位到对应输入框；当用户点击重置按钮时，页面会弹出确认框，确认后才真正清空表单。相关代码如图 3 所示。")
    add_picture(document, ASSETS / "code-form.png", "图 3 登录表单校验代码截图", 15.5)
    add_body_paragraph(document, "核心校验代码如下：")
    add_code_block(document, """function validateLoginForm() {
  const username = usernameInput.value.trim();
  const password = passwordInput.value.trim();

  if (!username) {
    usernameMessage = "用户名不能为空。";
  } else if (username.length < 8 || username.length > 20) {
    usernameMessage = "用户名长度应在 8~20 个字符之间。";
  }

  if (!password) {
    passwordMessage = "密码不能为空。";
  } else if (password.length < 6 || password.length > 8) {
    passwordMessage = "密码长度应在 6~8 个字符之间。";
  }
}""")

    add_heading(document, "五、实验结果")
    add_body_paragraph(document, "本实验运行环境为 macOS（Apple M1 Pro），使用 Visual Studio Code 编写代码，使用 Google Chrome 打开本地页面进行运行和测试。实验二源码位于实验二目录下的 index.html 文件。")
    steps = [
        "新建实验二页面文件，完成 HTML 结构、meta 标签和双栏布局设计。",
        "实现字号切换功能，并通过三种事件绑定方式演示单击事件处理。",
        "添加鼠标移动、键盘输入和窗口加载状态显示模块。",
        "设计登录表单，并实现用户名和密码的合法性校验。",
        "通过浏览器测试字号切换、中字号效果和非法登录提示效果。"
    ]
    for i, item in enumerate(steps, start=1):
        add_body_paragraph(document, f"{i}. {item}")
    add_body_paragraph(document, "图 4 为页面初始状态效果图，图 5 为点击“中”字号链接后的效果图，图 6 为输入不合法时的登录校验结果图。")
    add_picture(document, ASSETS / "page-initial.png", "图 4 实验二页面初始状态", 15.8)
    add_picture(document, ASSETS / "page-medium.png", "图 5 点击中字号后的页面效果", 15.8)
    add_picture(document, ASSETS / "page-login-error.png", "图 6 登录输入不合法时的提示效果", 15.8)
    add_body_paragraph(document, "测试结果表明，页面能够正确完成字号切换、事件状态更新、表单输入校验、错误提示显示和重置确认等功能，满足实验指导书提出的要求。")

    add_heading(document, "六、实验总结")
    add_body_paragraph(document, "通过本次实验，我进一步理解了 JavaScript 事件驱动机制以及前端表单校验的基本实现方法。与实验一相比，本次实验更加注重交互逻辑设计，不仅需要完成页面样式和结构的搭建，还需要根据用户输入和操作实时调整页面状态。")
    add_body_paragraph(document, "实验过程中，重点在于对用户名和密码校验规则的梳理，以及提交、重置、焦点、键盘和鼠标等事件的配合使用。为保证页面演示效果，我将两个项目整合在同一个页面中，并增加状态面板展示事件反馈，使实验内容更直观、也更方便截图和汇报。")
    add_body_paragraph(document, "总体来看，本次实验顺利完成了字号切换和表单校验两个任务，同时也加深了我对 HTML5 表单元素、JavaScript DOM 操作和常见浏览器事件的理解，为后续移动互联网开发课程中的交互式页面设计打下了基础。")

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
