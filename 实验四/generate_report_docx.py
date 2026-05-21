from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
ASSETS = BASE_DIR / "实验素材"
OUTPUT = BASE_DIR / "实验文档" / "2315302125 崔子霖4.docx"


def set_run_font(run, size=12, bold=False, east_asia="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_heading(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True, east_asia="黑体")


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, east_asia="宋体", western="Courier New")


def add_picture(document, image_path, caption, width_cm):
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = document.add_paragraph()
    format_paragraph(caption_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=11)


def build_cover(document):
    for _ in range(4):
        document.add_paragraph()

    school = document.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("南京理工大学紫金学院")
    set_run_font(run, size=22, bold=True, east_asia="黑体")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(42)
    run = title.add_run("《移动互联网开发基础》实验报告")
    set_run_font(run, size=20, bold=True, east_asia="黑体")

    table = document.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(8.8)

    fields = [
        ("课程编号：", "待补充"),
        ("课程名称：", "移动互联网开发基础"),
        ("学    院：", "计算机与人工智能"),
        ("专    业：", "计算机科学与技术"),
        ("学    号：", "2315302125"),
        ("姓    名：", "崔子霖"),
        ("任课教师：", "待补充"),
    ]

    for row, (label, value) in zip(table.rows, fields):
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        r2 = p2.add_run(value)
        set_run_font(r1, size=14)
        set_run_font(r2, size=14)

    for _ in range(8):
        document.add_paragraph()

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026 年 5 月")
    set_run_font(run, size=14)
    document.add_page_break()


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    build_cover(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("实验四 Html5 设计展邀请函")
    set_run_font(run, size=16, bold=True, east_asia="黑体")

    add_heading(document, "一、实验简介")
    add_body_paragraph(document, "本实验围绕 Swiper 插件的移动端页面滑动效果展开，最终完成了一份 Html5 设计展 H5 邀请函。页面采用三屏竖向滑动结构，分别展示活动封面、活动介绍和创意作品展示，并为不同页面元素设置进入动画。")
    add_body_paragraph(document, "实验环境为 macOS，开发工具为 Visual Studio Code，运行浏览器为 Google Chrome。源码中将 HTML 结构、CSS 样式和 JavaScript 交互逻辑分离，图片素材统一放入 img 文件夹，便于维护和提交。")

    add_heading(document, "二、实验要求")
    for index, item in enumerate([
        "掌握 Swiper 插件的使用方法，实现网页内容的滑动效果。",
        "理解 Swiper 插件的配置参数及其作用。",
        "学会在网页中嵌入 Swiper 插件，并实现垂直滑动效果。",
        "掌握 Swiper 动画插件的使用，为滑动内容添加动画效果。",
        "使用 Swiper Animate 设计一份三页邀请函，页面可在竖直方向上滑动切换。",
        "第一页标题使用 swing 效果，第二页标题使用 bounceIn 效果，第三页图片 pic1 使用 zoomIn 效果。",
    ], start=1):
        add_body_paragraph(document, f"{index}. {item}")

    add_heading(document, "三、概要设计")
    add_body_paragraph(document, "邀请函整体采用移动端 H5 页面结构，外层为 .swiper 容器，中间为 .swiper-wrapper，内部三页内容均使用 .swiper-slide 表示。第一页展示活动标题和艺术作品图，第二页展示活动介绍、地点和时间，第三页按照实验要求重点展示 pic1 图片。")
    add_body_paragraph(document, "页面通过 Swiper 的 direction 配置设置为 vertical，实现竖向滑动；通过 pagination 显示分页器；通过 mousewheel 配置支持桌面浏览器滚轮切换。动画部分通过元素上的 swiper-animate-effect、swiper-animate-duration 和 swiper-animate-delay 属性控制。")

    add_heading(document, "四、详细设计")
    add_body_paragraph(document, "HTML 部分定义 Swiper 基本结构，并为三个页面分别设置封面、介绍和作品展示内容。第一页标题使用 swing，第二页标题使用 bounceIn，第三页图片 pic1 使用 zoomIn。HTML 页面结构代码截图如图 1 所示。")
    add_picture(document, ASSETS / "code-html.png", "图 1 HTML 页面结构与动画属性代码截图", 15.5)
    add_body_paragraph(document, "CSS 部分负责移动端画布、背景图片、标题文字、作品图片、分页器和动画关键帧样式。JavaScript 部分负责初始化 Swiper，并在初始化和滑动结束时触发当前页面的动画。CSS 与 JavaScript 核心代码截图如图 2 所示。")
    add_picture(document, ASSETS / "code-css-js.png", "图 2 CSS 动画与 Swiper 初始化代码截图", 15.5)
    add_body_paragraph(document, "Swiper 初始化核心代码如下：")
    add_code_block(document, """var swiper = new Swiper(".invite-swiper", {
  direction: "vertical",
  loop: false,
  speed: 620,
  mousewheel: true,
  pagination: {
    el: ".swiper-pagination",
    clickable: true
  }
});""")

    add_heading(document, "五、实验结果")
    add_body_paragraph(document, "完成后的实验四源码位于实验四目录下，其中 index.html 为页面入口，css/style.css 为样式文件，js/main.js 为交互脚本，img 目录保存邀请函使用的图片素材。使用 Google Chrome 打开本地页面后，三页内容均能正常显示，并可通过竖向滑动或分页器切换。")
    for index, item in enumerate([
        "创建实验四目录，并建立 index.html、css、js、img、实验文档和实验素材等目录。",
        "按照 Swiper 结构编写三页邀请函页面。",
        "配置 Swiper 竖向滑动、分页器和鼠标滚轮切换。",
        "为第一页标题、第二页标题和第三页 pic1 图片分别设置 swing、bounceIn 和 zoomIn 动画。",
        "使用 Chrome 对三页效果和代码内容进行截图，并将截图整理进实验报告。",
    ], start=1):
        add_body_paragraph(document, f"{index}. {item}")
    add_picture(document, ASSETS / "page-slide-1.png", "图 3 邀请函第一页封面效果图", 7.2)
    add_picture(document, ASSETS / "page-slide-2.png", "图 4 邀请函第二页活动介绍效果图", 7.2)
    add_picture(document, ASSETS / "page-slide-3.png", "图 5 邀请函第三页作品展示效果图", 7.2)
    add_body_paragraph(document, "测试结果表明，页面结构完整，图片素材能够正常加载，三页内容可以竖向切换，动画效果能够在页面显示时触发，满足实验指导书中对 Swiper 竖向滑动和动画效果的要求。")

    add_heading(document, "六、实验总结")
    add_body_paragraph(document, "通过本次实验，我进一步熟悉了 Swiper 插件在移动端 H5 页面中的使用方式。与普通静态页面相比，Swiper 页面更强调屏幕之间的切换关系，需要将每一屏内容作为独立 slide 进行组织。")
    add_body_paragraph(document, "实验过程中，重点是正确搭建 Swiper 的 HTML 层级结构，并在 JavaScript 中配置竖向滑动、分页器和滑动回调。动画部分通过统一的 ani 类和动画属性控制，使不同页面可以使用不同进入效果。")
    add_body_paragraph(document, "总体来看，本次实验完成了 Html5 设计展邀请函的设计与实现，掌握了 Swiper 插件、移动端 H5 页面布局和页面元素动画的综合使用方法。")

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
