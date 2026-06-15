from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
INPUT = BASE_DIR / "实验文档" / "实验四答辩知识点.md"
OUTPUT = BASE_DIR / "实验文档" / "实验四答辩知识点.docx"


def set_run_font(run, size=11, bold=False, east_asia="宋体", western="Times New Roman"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(22)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=False)
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        set_run_font(run, size=15, bold=True, east_asia="黑体")
        run.font.color.rgb = RGBColor(31, 70, 128)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(7)
        run = paragraph.add_run(text)
        set_run_font(run, size=13, bold=True, east_asia="黑体")
        run.font.color.rgb = RGBColor(36, 96, 75)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=12, bold=True, east_asia="黑体")


def add_body(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_quote(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=False)
    paragraph.paragraph_format.left_indent = Cm(0.6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, east_asia="楷体")
    run.font.color.rgb = RGBColor(70, 70, 70)


def add_code(document, lines):
    text = "\n".join(lines)
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=False)
    paragraph.paragraph_format.left_indent = Cm(0.4)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, east_asia="宋体", western="Courier New")


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=False)
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=11)


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, first_indent=False)
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.3)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_table(document, rows):
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_index == 0:
                shade_cell(cell, "DCEAF7")
            paragraph = cell.paragraphs[0]
            set_paragraph(paragraph, first_indent=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5, bold=(r_index == 0), east_asia="黑体" if r_index == 0 else "宋体")
    document.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        if set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        cells = [cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def clean_inline(text):
    return text.replace("`", "")


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验四答辩知识点")
    set_run_font(run, size=20, bold=True, east_asia="黑体")
    run.font.color.rgb = RGBColor(31, 70, 128)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Html5 设计展邀请函")
    set_run_font(run, size=13, bold=True, east_asia="黑体")

    lines = INPUT.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(document, parse_table(table_lines))
            table_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if in_code:
                add_code(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            add_heading(document, clean_inline(line[3:]), 1)
        elif line.startswith("### "):
            add_heading(document, clean_inline(line[4:]), 2)
        elif line.startswith("> "):
            add_quote(document, clean_inline(line[2:]))
        elif line.startswith("- "):
            add_bullet(document, clean_inline(line[2:]))
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            add_numbered(document, clean_inline(line))
        else:
            add_body(document, clean_inline(line))

    flush_table()
    document.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
